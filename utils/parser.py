# utils/parser.py

# -------------------------------
# 📦 Imports
# -------------------------------
import os
import re
import pdfplumber
from docx import Document
from dotenv import load_dotenv
from openai import OpenAI
from utils.openai_client import client

# Load environment variables
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# -------------------------------
# 🔗 Extract Links from Text
# -------------------------------
def extract_links(text):
    """
    Extract all URLs from a given text.
    """
    pattern = r'https?://\S+|www\.\S+'
    return re.findall(pattern, text)

# -------------------------------
# 📄 Extract Text from PDF
# -------------------------------
def extract_text_from_pdf(file):
    """
    Extract text content from a PDF file.
    """
    text = ""
    try:
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
    except Exception as e:
        print("PDF Error:", e)
    return text

# -------------------------------
# 📄 Extract Text from DOCX
# -------------------------------
def extract_text_from_docx(file):
    """
    Extract text content from a DOCX file.
    """
    text = ""
    try:
        doc = Document(file)
        text = "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        print("DOCX Error:", e)
    return text

# -------------------------------
# 🔍 Clean Text
# -------------------------------
def clean_text(text):
    """
    Basic text cleaning: lowercase, remove newlines, punctuation.
    """
    text = text.lower()
    text = re.sub(r'\n', ' ', text)
    text = re.sub(r'[^\w\s]', '', text)
    return text

# -------------------------------
# 🤖 Extract Skills using OpenAI (JD)
# -------------------------------
def extract_skills_with_ai(jd_text):
    """
    Extract top 5 technical skills from a Job Description using OpenAI.
    """
    if not jd_text.strip():
        return []

    prompt = f"""
Extract top 5 technical skills from the following job description.

Rules:
- Return ONLY a comma-separated list
- No explanation
- Max 5 skills

Job Description:
{jd_text}
"""
    try:
        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": "You are a technical skill extractor."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3
        )

        skills_text = response.choices[0].message.content.strip()
        skills = [skill.strip() for skill in skills_text.split(",")]
        return skills[:5]

    except Exception as e:
        print("OpenAI Error:", e)
        return []

# -------------------------------
# 🧠 Extract Skills from Resume using OpenAI
# -------------------------------
def extract_resume_skills_with_ai(resume_text):
    """
    Extract top 5 technical skills from a Resume using OpenAI.
    """
    if not resume_text.strip():
        return []

    prompt = f"""
Extract top 5 technical skills from this resume.

Rules:
- Return ONLY comma-separated skills
- No explanation

Resume:
{resume_text}
"""
    try:
        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": "You are a resume skill extractor."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3
        )

        skills_text = response.choices[0].message.content.strip()
        skills = [s.strip() for s in skills_text.split(",")]
        return skills[:5]

    except Exception as e:
        print("OpenAI Error:", e)
        return []

# -------------------------------
# 📊 Match Resume Skills vs JD Skills
# -------------------------------
def match_skills(resume_skills, jd_skills):
    """
    Compare resume skills with JD skills and calculate match percentage.
    """
    resume_set = set([s.lower() for s in resume_skills])
    jd_set = set([s.lower() for s in jd_skills])

    matched = list(resume_set & jd_set)
    missing = list(jd_set - resume_set)

    match_percent = (len(matched) / len(jd_set)) * 100 if jd_set else 0

    return {
        "matched": matched,
        "missing": missing,
        "match_percent": round(match_percent, 2)
    }

# -------------------------------
# 🔄 Alternative JD Skill Extractor
# -------------------------------
def extract_jd_skills(jd_text):
    """
    Quick helper to extract skills from JD (returns list).
    """
    prompt = f"""
Extract top 5 technical skills from this job description.
Return only comma separated skills.

JD:
{jd_text}
"""
    try:
        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3
        )
        skills = [s.strip() for s in response.choices[0].message.content.strip().split(",")]
        return skills
    except Exception as e:
        print("OpenAI Error:", e)
        return []
    
    # resume_utils.py

# -------------------------------
# 📦 Imports
# -------------------------------
import os
from docx import Document
from docx.shared import Pt
from utils.ai_bullet_generator import generate_experience_bullets

# -------------------------------
# 📝 Helper: Add bullet paragraph safely
# -------------------------------
def add_bullet_paragraph(doc, text):
    """
    Adds a bullet point safely, avoiding style errors.
    """
    para = doc.add_paragraph(f"• {text}")  # manual bullet
    # Set font size manually
    for run in para.runs:
        run.font.size = Pt(11)

# -------------------------------
# 📝 Helper: Add skills table
# -------------------------------
def add_skills_table(doc, skills, title="Skills"):
    """
    Adds a table of skills.
    """
    if not skills:
        return
    doc.add_paragraph(f"{title}:")
    table = doc.add_table(rows=1, cols=len(skills))
    table.autofit = True
    hdr_cells = table.rows[0].cells
    for i, skill in enumerate(skills):
        hdr_cells[i].text = skill.capitalize()
    doc.add_paragraph("\n")

# -------------------------------
# 📄 Update DOCX Resume
# -------------------------------
def update_docx_resume(file_path, skills=None, linkedin=None, github=None):
    """
    Updates DOCX resume:
    - Adds missing JD skills (table)
    - Appends AI-generated bullets
    - Optionally adds LinkedIn/GitHub links
    """
    doc = Document(file_path)

    # Extract existing text
    resume_text = "\n".join([p.text for p in doc.paragraphs])

    # Generate AI experience bullets
    ai_bullets = generate_experience_bullets(resume_text, skills or [])
    if ai_bullets:
        doc.add_paragraph("\nExperience Highlights:")
        for bullet in ai_bullets.split("\n"):
            add_bullet_paragraph(doc, bullet)

    # Add skills table
    if skills:
        add_skills_table(doc, skills, title="Skills")

    # Add LinkedIn/GitHub links
    if linkedin:
        doc.add_paragraph(f"LinkedIn: {linkedin}")
    if github:
        doc.add_paragraph(f"GitHub: {github}")

    # Save updated resume
    output_path = os.path.splitext(file_path)[0] + "_updated.docx"
    doc.save(output_path)

    return output_path

# -------------------------------
# 📄 Placeholder: Update PDF Resume
# -------------------------------
def update_pdf_resume(file_path, skills=None, linkedin=None, github=None):
    """
    Currently returns original path. Implement PDF logic separately.
    """
    return file_path