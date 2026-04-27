# utils/ats_formatter.py

import re
import random
from io import BytesIO

import pdfplumber
from docx import Document
from openai import OpenAI
from PyPDF2 import PdfReader, PdfWriter
from reportlab.pdfgen import canvas

client = OpenAI()

# -----------------------------------
# 🤖 Extract JD Skills using OpenAI
# -----------------------------------
def extract_jd_skills(jd_text):
    prompt = f"""
Extract top 5 technical skills from this job description.

Rules:
- Return ONLY comma-separated skills
- No explanation

Job Description:
{jd_text}
"""
    try:
        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3
        )
        skills = response.choices[0].message.content.strip()
        return [s.strip() for s in skills.split(",")]
    except Exception as e:
        print("OpenAI Error:", e)
        return []

# -----------------------------------
# 📝 Generate experience snippet for a skill
# -----------------------------------
def generate_experience_for_skill(skill):
    templates = [
        f"Implemented {skill} to improve project efficiency and reduce errors.",
        f"Applied {skill} in key projects to deliver high-quality results.",
        f"Leveraged {skill} to streamline operations and optimize workflows.",
        f"Utilized {skill} to develop scalable and robust solutions."
    ]
    return random.choice(templates)

# -----------------------------------
# 📄 Update DOCX (PRESERVE FORMAT)
# -----------------------------------
def update_docx_add_skills(file_path, jd_text, output_file="output/ATS_Resume.docx"):
    doc = Document(file_path)
    jd_skills = extract_jd_skills(jd_text)

    inserted = False
    for para in doc.paragraphs:
        text = para.text.lower()
        if "skills" in text:
            existing_text = para.text
            new_skills = [skill for skill in jd_skills if skill.lower() not in existing_text.lower()]
            if new_skills:
                para.add_run(", " + ", ".join(new_skills))
            inserted = True
            break

    if not inserted and jd_skills:
        doc.add_paragraph("\nSkills: " + ", ".join(jd_skills))

    doc.save(output_file)
    return output_file

def update_resume(doc, existing_skills, new_skills, skills_style):
    # Only add skills not already present
    to_add = [s for s in new_skills if s.lower() not in [x.lower() for x in existing_skills]]

    if skills_style == 'table':
        # Add to first table containing "skills"
        skills_table = None
        for table in doc.tables:
            for row in table.rows:
                if 'skills' in row.cells[0].text.lower():
                    skills_table = table
                    break
            if skills_table:
                break
        if skills_table:
            for skill in to_add:
                row = skills_table.add_row()
                row.cells[0].text = skill

    elif skills_style == 'bullet':
        # Add as bullet points after existing skills
        for para in doc.paragraphs:
            if para.style.name.startswith('List') and 'skill' in para.text.lower():
                for skill in to_add:
                    doc.add_paragraph(skill, style=para.style)
                break

    # Generate experience entries for new skills
    for skill in to_add:
        exp_desc = generate_experience_for_skill(skill)
        doc.add_paragraph(exp_desc, style='Normal')

    return doc

# -----------------------------------
# 📄 Extract PDF text
# -----------------------------------
def extract_pdf_text(file_path):
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text

# -----------------------------------
# 📄 Update PDF (LIMITED SAFE MODE)
# -----------------------------------
def update_pdf_add_skills(file_path, jd_text, output_file="output/ATS_Resume.pdf"):
    reader = PdfReader(file_path)
    writer = PdfWriter()
    jd_skills = extract_jd_skills(jd_text)

    # Create overlay with skills
    packet = BytesIO()
    can = canvas.Canvas(packet)
    text = "Added Skills: " + ", ".join(jd_skills)
    can.drawString(50, 50, text)  # bottom of page
    can.save()
    packet.seek(0)
    overlay = PdfReader(packet)

    for page in reader.pages:
        page.merge_page(overlay.pages[0])
        writer.add_page(page)

    with open(output_file, "wb") as f:
        writer.write(f)

    return output_file