# resume_utils.py

import os
from docx import Document
from utils.ai_bullet_generator import generate_experience_bullets

def add_manual_bullet(doc, text):
    """
    Adds a bullet paragraph manually to avoid 'List Bullet' style errors.
    Falls back to a normal paragraph with a bullet if needed.
    """
    try:
        doc.add_paragraph(text, style='List Bullet')
    except Exception:
        doc.add_paragraph(f"• {text}")

def add_skills_table(doc, skills, title="Skills"):
    """
    Adds a table of skills to the document.
    """
    if not skills:
        return
    doc.add_paragraph(f"{title}:")
    table = doc.add_table(rows=1, cols=len(skills))
    hdr_cells = table.rows[0].cells
    for i, skill in enumerate(skills):
        hdr_cells[i].text = skill.capitalize()
    doc.add_paragraph("\n")  # spacing after table

def update_docx_resume(file_path, skills=None, linkedin=None, github=None):
    """
    Updates a DOCX resume:
    - Adds missing JD skills (table)
    - Appends AI-generated experience bullets
    - Adds LinkedIn/GitHub links
    """
    doc = Document(file_path)

    # Extract existing resume text
    resume_text = "\n".join([p.text for p in doc.paragraphs])

    # Generate AI bullets safely
    try:
        ai_bullets = generate_experience_bullets(resume_text, skills or [])
    except Exception as e:
        print("OpenAI API failed:", e)
        ai_bullets = (
            "• Improved system performance.\n"
            "• Optimized existing processes.\n"
            "• Implemented scalable solutions."
        )

    # Add experience bullets
    if ai_bullets:
        doc.add_paragraph("Experience Highlights:")
        for bullet in ai_bullets.split("\n"):
            bullet = bullet.strip()
            if bullet:
                add_manual_bullet(doc, bullet)

    # Add skills table
    if skills:
        add_skills_table(doc, skills, title="Skills")

    # Add LinkedIn/GitHub links
    if linkedin:
        doc.add_paragraph(f"LinkedIn: {linkedin}")
    if github:
        doc.add_paragraph(f"GitHub: {github}")

    # Save updated resume
    output_path = os.path.splitext(file_path)[0] + "_updated.docx"
    doc.save(output_path)
    return output_path

def update_pdf_resume(file_path, skills=None, linkedin=None, github=None):
    """
    Placeholder: return original PDF for now.
    PDF update (skills/bullets) requires separate implementation.
    """
    return file_path